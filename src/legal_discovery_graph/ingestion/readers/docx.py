"""DOCX reader: extract plain text from a Word document into a ``Document``/body pair."""

from pathlib import Path

import docx as _python_docx
from docx.opc.exceptions import PackageNotFoundError

from legal_discovery_graph.ingestion.readers._common import file_document_id, normalize_paragraphs
from legal_discovery_graph.models import Document, DocumentType


def read_docx(path: Path, custodian: str = "") -> tuple[Document, str]:
    """Read a DOCX file, returning its ``Document`` metadata and extracted plain-text body.

    Raises:
        ValueError: if the file cannot be opened as a DOCX or contains no text.
    """
    try:
        word_document = _python_docx.Document(str(path))
    except (PackageNotFoundError, KeyError, ValueError) as exc:
        raise ValueError(f"Unable to read DOCX '{path}': {exc}") from exc

    paragraphs = [paragraph.text for paragraph in word_document.paragraphs]
    body = normalize_paragraphs(paragraphs)
    if not body:
        raise ValueError(f"No extractable text found in DOCX '{path}'")

    document = Document(
        document_id=file_document_id(path),
        doc_type=DocumentType.OTHER,
        title=path.stem,
        source_path=str(path),
        custodian=custodian,
    )
    return document, body
