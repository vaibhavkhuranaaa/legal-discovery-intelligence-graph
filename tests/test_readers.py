"""Tests for the pdf/docx/eml file-format readers in `ingestion.readers`."""

from datetime import UTC
from email.message import EmailMessage
from pathlib import Path

import docx as python_docx
import pytest

from legal_discovery_graph.ingestion.readers import read_docx, read_eml, read_pdf
from legal_discovery_graph.models import DocumentType


def _make_minimal_pdf_bytes(text: str) -> bytes:
    """Build a minimal one-page PDF with a text-showing content stream."""
    content = f"BT /F1 12 Tf 10 100 Td ({text}) Tj ET".encode()
    header = b"%PDF-1.4\n"
    objs = [
        b"1 0 obj\n<</Type/Catalog/Pages 2 0 R>>\nendobj\n",
        b"2 0 obj\n<</Type/Pages/Kids[3 0 R]/Count 1>>\nendobj\n",
        b"3 0 obj\n<</Type/Page/Parent 2 0 R/Resources<</Font<</F1 4 0 R>>>>"
        b"/MediaBox[0 0 200 200]/Contents 5 0 R>>\nendobj\n",
        b"4 0 obj\n<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>\nendobj\n",
        b"5 0 obj\n<</Length "
        + str(len(content)).encode()
        + b">>\nstream\n"
        + content
        + b"\nendstream\nendobj\n",
    ]
    offsets = []
    pos = len(header)
    body = b""
    for obj in objs:
        offsets.append(pos)
        body += obj
        pos += len(obj)

    xref_offset = len(header) + len(body)
    xref = b"xref\n0 6\n0000000000 65535 f \n"
    for offset in offsets:
        xref += f"{offset:010d} 00000 n \n".encode()
    trailer = (
        b"trailer\n<</Size 6/Root 1 0 R>>\nstartxref\n" + str(xref_offset).encode() + b"\n%%EOF"
    )
    return header + body + xref + trailer


def _make_blank_pdf_bytes() -> bytes:
    """Build a minimal one-page PDF with no content stream (no extractable text)."""
    header = b"%PDF-1.4\n"
    objs = [
        b"1 0 obj\n<</Type/Catalog/Pages 2 0 R>>\nendobj\n",
        b"2 0 obj\n<</Type/Pages/Kids[3 0 R]/Count 1>>\nendobj\n",
        b"3 0 obj\n<</Type/Page/Parent 2 0 R/MediaBox[0 0 200 200]>>\nendobj\n",
    ]
    offsets = []
    pos = len(header)
    body = b""
    for obj in objs:
        offsets.append(pos)
        body += obj
        pos += len(obj)

    xref_offset = len(header) + len(body)
    xref = b"xref\n0 4\n0000000000 65535 f \n"
    for offset in offsets:
        xref += f"{offset:010d} 00000 n \n".encode()
    trailer = (
        b"trailer\n<</Size 4/Root 1 0 R>>\nstartxref\n" + str(xref_offset).encode() + b"\n%%EOF"
    )
    return header + body + xref + trailer


class TestReadPdf:
    def test_happy_path_extracts_text_and_metadata(self, tmp_path: Path) -> None:
        pdf_path = tmp_path / "falcon_memo.pdf"
        pdf_path.write_bytes(_make_minimal_pdf_bytes("Falcon payment approved"))

        document, body = read_pdf(pdf_path, custodian="alice")

        assert "Falcon payment approved" in body
        assert document.doc_type == DocumentType.OTHER
        assert document.title == "falcon_memo"
        assert document.source_path == str(pdf_path)
        assert document.custodian == "alice"

    def test_empty_extraction_raises_value_error(self, tmp_path: Path) -> None:
        pdf_path = tmp_path / "blank.pdf"
        pdf_path.write_bytes(_make_blank_pdf_bytes())

        with pytest.raises(ValueError, match="No extractable text"):
            read_pdf(pdf_path)

    def test_document_id_is_repeatable(self, tmp_path: Path) -> None:
        pdf_path = tmp_path / "falcon_memo.pdf"
        pdf_path.write_bytes(_make_minimal_pdf_bytes("Falcon payment approved"))

        document_1, _ = read_pdf(pdf_path)
        document_2, _ = read_pdf(pdf_path)

        assert document_1.document_id == document_2.document_id


class TestReadDocx:
    def test_happy_path_joins_paragraphs_with_blank_line(self, tmp_path: Path) -> None:
        docx_path = tmp_path / "notes.docx"
        word_document = python_docx.Document()
        word_document.add_paragraph("Falcon project kickoff meeting.")
        word_document.add_paragraph("Budget approved by Jane Doe.")
        word_document.save(str(docx_path))

        document, body = read_docx(docx_path, custodian="bob")

        assert body == ("Falcon project kickoff meeting.\n\nBudget approved by Jane Doe.")
        assert document.doc_type == DocumentType.OTHER
        assert document.title == "notes"
        assert document.custodian == "bob"

    def test_document_id_is_repeatable_across_reads(self, tmp_path: Path) -> None:
        docx_path = tmp_path / "notes.docx"
        word_document = python_docx.Document()
        word_document.add_paragraph("Some content.")
        word_document.save(str(docx_path))

        document_1, _ = read_docx(docx_path)
        document_2, _ = read_docx(docx_path)

        assert document_1.document_id == document_2.document_id

    def test_empty_document_raises_value_error(self, tmp_path: Path) -> None:
        docx_path = tmp_path / "empty.docx"
        word_document = python_docx.Document()
        word_document.save(str(docx_path))

        with pytest.raises(ValueError, match="No extractable text"):
            read_docx(docx_path)


def _write_eml(tmp_path: Path, message: EmailMessage, filename: str = "message.eml") -> Path:
    eml_path = tmp_path / filename
    eml_path.write_bytes(bytes(message))
    return eml_path


class TestReadEml:
    def test_multipart_prefers_text_plain_and_parses_headers(self, tmp_path: Path) -> None:
        message = EmailMessage()
        message["From"] = "alice@example.com"
        message["To"] = "bob@example.com"
        message["Subject"] = "Falcon Contract Review"
        message["Date"] = "Tue, 14 Jul 2026 10:30:00 -0400"
        message.set_content("Plain text body about the Falcon contract.")
        message.add_alternative(
            "<html><body><p>HTML body about the Falcon contract.</p></body></html>",
            subtype="html",
        )

        eml_path = _write_eml(tmp_path, message)
        document, body = read_eml(eml_path, custodian="carol")

        assert document.title == "Falcon Contract Review"
        assert document.doc_type == DocumentType.EMAIL
        assert document.custodian == "carol"
        assert document.sent_at is not None
        assert document.sent_at.tzinfo is not None

        header_block, _, rest = body.partition("\n\n")
        assert "From: alice@example.com" in header_block
        assert "To: bob@example.com" in header_block
        assert "Subject: Falcon Contract Review" in header_block
        assert "Date:" in header_block
        assert "Plain text body about the Falcon contract." in rest
        assert "<p>" not in rest

    def test_html_only_falls_back_and_strips_tags(self, tmp_path: Path) -> None:
        message = EmailMessage()
        message["From"] = "alice@example.com"
        message["To"] = "bob@example.com"
        message["Subject"] = "HTML Only Notice"
        message["Date"] = "Tue, 14 Jul 2026 10:30:00 -0400"
        message.set_content(
            "<html><body><p>Only HTML content here.</p></body></html>",
            subtype="html",
        )

        eml_path = _write_eml(tmp_path, message, filename="html_only.eml")
        _, body = read_eml(eml_path)

        assert "Only HTML content here." in body
        assert "<p>" not in body
        assert "<html>" not in body

    def test_missing_date_header_leaves_sent_at_none(self, tmp_path: Path) -> None:
        message = EmailMessage()
        message["From"] = "alice@example.com"
        message["To"] = "bob@example.com"
        message["Subject"] = "No Date Header"
        message.set_content("Body with no date header.")

        eml_path = _write_eml(tmp_path, message, filename="no_date.eml")
        document, body = read_eml(eml_path)

        assert document.sent_at is None
        assert "Date:" not in body.partition("\n\n")[0]

    def test_subject_missing_falls_back_to_filename_stem(self, tmp_path: Path) -> None:
        message = EmailMessage()
        message["From"] = "alice@example.com"
        message["To"] = "bob@example.com"
        message.set_content("Body with no subject.")

        eml_path = _write_eml(tmp_path, message, filename="no_subject.eml")
        document, _ = read_eml(eml_path)

        assert document.title == "no_subject"

    def test_document_id_is_repeatable(self, tmp_path: Path) -> None:
        message = EmailMessage()
        message["From"] = "alice@example.com"
        message["To"] = "bob@example.com"
        message["Subject"] = "Repeatable"
        message["Date"] = "Tue, 14 Jul 2026 10:30:00 -0400"
        message.set_content("Body content.")

        eml_path = _write_eml(tmp_path, message, filename="repeatable.eml")
        document_1, _ = read_eml(eml_path)
        document_2, _ = read_eml(eml_path)

        assert document_1.document_id == document_2.document_id

    def test_naive_date_gets_utc_timezone(self) -> None:
        from email import message_from_string

        from legal_discovery_graph.ingestion.readers.eml import _parse_sent_at

        message = message_from_string("Date: Tue, 14 Jul 2026 10:30:00\n\nBody")
        sent_at = _parse_sent_at(message)

        assert sent_at is not None
        assert sent_at.tzinfo is UTC
